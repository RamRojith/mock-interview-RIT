from django.core.management.base import BaseCommand
from django.db import connection
from django.conf import settings
from docx import Document
import json
import os
from collections import defaultdict
from datetime import datetime


class Command(BaseCommand):
    help = "Optimized DB audit to DOCX (append mode)"

    def handle(self, *args, **kwargs):
        cursor = connection.cursor()

        # 🔥 Only new logs
        cursor.execute("""
            SELECT id, table_name, action_type, record_id, old_data, new_data, db_user, changed_at
            FROM audit_raw_logs
            WHERE processed = 0
            ORDER BY changed_at ASC
        """)

        rows = cursor.fetchall()

        if not rows:
            self.stdout.write("No new logs")
            return

        # 📦 Group logs by date (important optimization)
        grouped_logs = defaultdict(list)

        for row in rows:
            log_id, table, action, record_id, old_data, new_data, db_user, changed_at = row

            dt = changed_at if isinstance(changed_at, datetime) else datetime.now()
            key = dt.strftime("%Y-%m-%d")

            grouped_logs[key].append({
                "id": log_id,
                "table": table,
                "action": action,
                "record_id": record_id,
                "old": json.loads(old_data) if old_data else {},
                "new": json.loads(new_data) if new_data else {},
                "user": db_user,
                "time": dt
            })

        # 📄 Process each date file once (FAST)
        for date_key, logs in grouped_logs.items():
            year, month, day = date_key.split("-")

            dir_path = os.path.join(settings.MEDIA_ROOT, "DB_logs", year, month)
            os.makedirs(dir_path, exist_ok=True)

            file_path = os.path.join(dir_path, f"{day}.docx")

            # Load or create document
            if os.path.exists(file_path):
                document = Document(file_path)
            else:
                document = Document()
                document.add_heading(f"Database Audit Logs - {date_key}", 0)

            # ✍️ Write logs
            for log in logs:
                document.add_paragraph(f"Time: {log['time']}")
                document.add_paragraph(f"User: {log['user']}")
                document.add_paragraph(f"Table: {log['table']}")
                document.add_paragraph(f"Action: {log['action']}")
                document.add_paragraph(f"Record ID: {log['record_id']}")

                if log["action"] == "UPDATE":
                    document.add_paragraph("OLD DATA:")
                    for k, v in log["old"].items():
                        document.add_paragraph(f"  {k}: {v}")

                    document.add_paragraph("NEW DATA:")
                    for k, v in log["new"].items():
                        document.add_paragraph(f"  {k}: {v}")

                elif log["action"] == "INSERT":
                    document.add_paragraph("NEW DATA:")
                    for k, v in log["new"].items():
                        document.add_paragraph(f"  {k}: {v}")

                elif log["action"] == "DELETE":
                    document.add_paragraph("DELETED DATA:")
                    for k, v in log["old"].items():
                        document.add_paragraph(f"  {k}: {v}")

                document.add_paragraph("-" * 60)

            document.save(file_path)

        # ✅ Mark all processed in ONE query (fast)
        ids = [str(r[0]) for r in rows]
        cursor.execute(
            f"UPDATE audit_raw_logs SET processed = 1 WHERE id IN ({','.join(ids)})"
        )

        self.stdout.write(self.style.SUCCESS("DOCX updated successfully (optimized)"))